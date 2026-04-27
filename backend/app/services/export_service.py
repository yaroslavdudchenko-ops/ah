import io
import logging
import mistune
from typing import Literal
from app.models.protocol import Protocol, ProtocolVersion

logger = logging.getLogger(__name__)

SECTION_TITLES = {
    "title_page": "Title Page / Титульная страница",
    "synopsis": "Synopsis / Краткое резюме",
    "introduction": "Introduction & Background / Введение",
    "objectives": "Study Objectives / Цели исследования",
    "design": "Study Design / Дизайн исследования",
    "population": "Study Population / Популяция",
    "treatment": "Study Treatment / Лечение",
    "efficacy": "Efficacy Assessments / Оценка эффективности",
    "safety": "Safety Assessments / Оценка безопасности",
    "statistics": "Statistical Analysis / Статистический анализ",
    "ethics": "Ethics / Этические аспекты",
    "references": "References / Список литературы",
}

ExportFormat = Literal["md", "html", "docx"]


def _build_markdown(protocol: Protocol, version: ProtocolVersion) -> str:
    content = version.content or {}
    lines = [
        f"# {protocol.title}",
        f"> **FOR DEMONSTRATION PURPOSES ONLY — SYNTHETIC DATA** | AI-Assisted. Requires qualified person review. | Версия: {version.version_number} | Модель: {version.generated_by}",
        "",
        "## Параметры исследования",
        f"- **Препарат:** {protocol.drug_name} ({protocol.inn})",
        f"- **Фаза:** {protocol.phase}",
        f"- **Терапевтическая область:** {protocol.therapeutic_area}",
        f"- **Индикация:** {protocol.indication}",
        f"- **Первичная конечная точка:** {protocol.primary_endpoint}",
        f"- **Длительность:** {protocol.duration_weeks} недель",
        f"- **Дозирование:** {protocol.dosing}",
        "",
    ]
    for section_key, text in content.items():
        title = SECTION_TITLES.get(section_key, section_key.replace("_", " ").title())
        lines.append(f"## {title}")
        lines.append(text)
        lines.append("")

    open_issues = getattr(protocol, "open_issues", [])
    if open_issues:
        lines.append("## Список открытых вопросов")
        for issue in open_issues:
            lines.append(f"- **[{issue.severity.upper()}]** {issue.section}: {issue.description}")
        lines.append("")

    return "\n".join(lines)


def export_markdown(protocol: Protocol, version: ProtocolVersion) -> bytes:
    return _build_markdown(protocol, version).encode("utf-8")


def export_html(protocol: Protocol, version: ProtocolVersion) -> bytes:
    md_text = _build_markdown(protocol, version)
    renderer = mistune.create_markdown()
    body = renderer(md_text)
    html = f"""<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{protocol.title}</title>
<style>
  :root {{
    --brand: #0284c7;
    --brand-dark: #0c4a6e;
    --brand-light: #f0f9ff;
    --text: #111827;
    --muted: #6b7280;
    --border: #e5e7eb;
  }}

  * {{ box-sizing: border-box; }}

  body {{
    font-family: 'Segoe UI', Arial, sans-serif;
    max-width: 900px;
    margin: 40px auto;
    padding: 0 28px;
    line-height: 1.7;
    color: var(--text);
    background: #fff;
  }}

  /* Document header */
  .doc-header {{
    border-bottom: 3px solid var(--brand);
    padding-bottom: 16px;
    margin-bottom: 32px;
  }}
  .doc-header h1 {{
    font-size: 26px;
    font-weight: 800;
    color: var(--brand-dark);
    margin: 0 0 6px;
    border: none;
    padding: 0;
  }}
  .doc-disclaimer {{
    background: var(--brand-light);
    border-left: 4px solid var(--brand);
    padding: 10px 16px;
    font-size: 12px;
    color: var(--muted);
    margin-top: 12px;
    border-radius: 0 6px 6px 0;
  }}

  h1 {{
    font-size: 22px;
    font-weight: 800;
    color: var(--brand-dark);
    border-bottom: 3px solid var(--brand);
    padding-bottom: 8px;
    margin: 0 0 24px;
  }}
  h2 {{
    font-size: 16px;
    font-weight: 700;
    color: var(--brand-dark);
    border-bottom: 2px solid var(--brand-light);
    margin-top: 36px;
    margin-bottom: 12px;
    padding-bottom: 6px;
    display: flex;
    align-items: center;
    gap: 8px;
  }}
  h2::before {{
    content: '';
    display: inline-block;
    width: 4px;
    height: 16px;
    background: var(--brand);
    border-radius: 2px;
    flex-shrink: 0;
  }}
  h3 {{
    font-size: 14px;
    font-weight: 600;
    color: #374151;
    margin-top: 20px;
    margin-bottom: 6px;
  }}

  p {{ margin: 0 0 10px; font-size: 14px; }}

  ul, ol {{ margin: 6px 0 12px 20px; }}
  li {{ font-size: 14px; line-height: 1.65; margin-bottom: 3px; }}

  strong {{ color: var(--brand-dark); }}

  /* Parameter block */
  .params-block {{
    background: var(--brand-light);
    border: 1px solid #bae6fd;
    border-radius: 8px;
    padding: 16px 20px;
    margin: 0 0 28px;
  }}
  .params-block ul {{ margin: 0; list-style: none; padding: 0; display: grid; grid-template-columns: 1fr 1fr; gap: 6px 24px; }}
  .params-block li {{ font-size: 13px; color: #374151; }}
  .params-block li strong {{ color: var(--brand-dark); }}

  blockquote {{
    background: var(--brand-light);
    border-left: 4px solid var(--brand);
    padding: 10px 16px;
    margin: 12px 0;
    border-radius: 0 6px 6px 0;
    font-size: 13px;
    color: #374151;
  }}

  table {{
    border-collapse: collapse;
    width: 100%;
    margin: 12px 0;
    font-size: 13px;
  }}
  th {{
    background: var(--brand);
    color: #fff;
    padding: 8px 12px;
    text-align: left;
    font-weight: 600;
  }}
  td {{ border: 1px solid var(--border); padding: 7px 12px; }}
  tr:nth-child(even) td {{ background: #f8faff; }}

  /* Open issues */
  .issue-critical {{ color: #dc2626; font-weight: 600; }}
  .issue-major    {{ color: #ea580c; font-weight: 600; }}
  .issue-minor    {{ color: #ca8a04; font-weight: 600; }}

  /* Footer */
  .doc-footer {{
    margin-top: 48px;
    padding-top: 12px;
    border-top: 2px solid var(--brand-light);
    font-size: 11px;
    color: var(--muted);
    text-align: center;
  }}

  /* Print */
  @media print {{
    @page {{ size: A4 portrait; margin: 20mm 15mm; }}
    * {{ -webkit-print-color-adjust: exact !important; print-color-adjust: exact !important; }}
    body {{ margin: 0; padding: 0; max-width: 100%; }}
    h2 {{ break-after: avoid; page-break-after: avoid; }}
    section {{ break-inside: avoid; page-break-inside: avoid; }}
  }}
</style>
</head>
<body>
{body}
<div class="doc-footer">Synthia AI Protocol Generator &nbsp;·&nbsp; FOR RESEARCH USE ONLY — NOT FOR CLINICAL USE</div>
</body>
</html>"""
    return html.encode("utf-8")


def export_docx(protocol: Protocol, version: ProtocolVersion) -> bytes:
    """P2 feature — DOCX export via python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading(protocol.title, level=0)
        doc.add_paragraph(
            "FOR DEMONSTRATION PURPOSES ONLY — SYNTHETIC DATA | "
            "AI-Assisted. Requires qualified person review. | "
            f"Version: {version.version_number} | Model: {version.generated_by}"
        )

        doc.add_heading("Параметры исследования", level=1)
        params = [
            f"Препарат: {protocol.drug_name} ({protocol.inn})",
            f"Фаза: {protocol.phase}",
            f"Терапевтическая область: {protocol.therapeutic_area}",
            f"Индикация: {protocol.indication}",
            f"Первичная конечная точка: {protocol.primary_endpoint}",
            f"Длительность: {protocol.duration_weeks} недель",
            f"Дозирование: {protocol.dosing}",
        ]
        for p in params:
            doc.add_paragraph(p, style="List Bullet")

        content = version.content or {}
        for section_key, text in content.items():
            title = SECTION_TITLES.get(section_key, section_key.replace("_", " ").title())
            doc.add_heading(title, level=1)
            for para in text.split("\n\n"):
                para = para.strip()
                if para.startswith("## ") or para.startswith("### "):
                    doc.add_heading(para.lstrip("#").strip(), level=2)
                elif para:
                    doc.add_paragraph(para)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as exc:
        logger.error("docx_export_failed", extra={"error": str(exc)})
        raise


CONTENT_TYPES: dict[ExportFormat, str] = {
    "md": "text/markdown; charset=utf-8",
    "html": "text/html; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

FILENAMES: dict[ExportFormat, str] = {
    "md": "protocol.md",
    "html": "protocol.html",
    "docx": "protocol.docx",
}
